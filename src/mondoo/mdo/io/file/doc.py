from __future__ import annotations

from mondoo.mdo.core.common     import (
    Paragraph,
    Figure,
    BaseModelWithBody
)

from .generic import FileDesc

from docx     import Document
from pathlib  import Path
from os       import PathLike
from typing   import Optional, List, Dict
from pydantic import Field
from PIL      import Image

import io

class Body:
    def __init__(
        self,
        paragraphs  : Optional[List[Paragraph]] = None,
        figures     : Optional[List[Figure]]    = None,
        total_words : int                = 0
    ):
        self.paragraphs  = paragraphs or []
        self.figures     = figures or []
        self.total_words = total_words
        
    def to_dict(self) -> Dict:
        return {
            "paragraphs" : [p.to_dict() for p in self.paragraphs],
            "figures"    : [i.to_dict() for i in self.figures]
        }
    
    def update(
        self,
        body : Body
    ) -> None:
        self.paragraphs    += body.paragraphs
        self.figures += body.figures
        self.total_words   += body.total_words


class DOCXObject(BaseModelWithBody):
    descriptor   : FileDesc    = Field(None, description="")
    body         : Body | None   = Field(None, description="")


def _extract_text_(path):
    doc = Document(path)
    text = []
    for para in doc.paragraphs:
        text.append(para.text.strip())

    full_text = "\n".join(text)
    print(full_text)
    

def _extract_illustrations_(doc):
    illustrations = []
    for i, rel in enumerate(doc.part._rels.values()):
        if 'image' not in rel.reltype:
            continue

        image_bytes = rel.target_part.blob
        image_ext   = rel.target_part.content_type.split("/")[-1]

        with Image.open(io.BytesIO(image_bytes)) as img:
            width, height = img.size

        illustration = Figure(
            page_ids    = [],
            image_bytes = image_bytes,
            image_ext   = image_ext,
            index       = i,
            width       = width,
            height      = height,
        )

        illustrations.append(illustration)
    
    return illustrations
                

def extract_docx_body(
    docx_path  : PathLike[str], 
    output_dir : Optional[PathLike[str] | None] = None
):
    doc = Document(docx_path)
    paragraphs = []
    total_count = 0
    for para in doc.paragraphs:
        text = para.text.strip()
        count = len(text)
        paragraphs.append(
            Paragraph(
                page_ids = [],
                content  = text,
                count    = count
            )
        )

        total_count += count
    
    illustrations = _extract_illustrations_(doc=doc)
     
    if output_dir is not None:
        output_dir = Path(output_dir)
        output_dir.mkdir(exist_ok=True)
        pages_out_path = output_dir / "page_data.json"
        image_dir      = output_dir / "images"
        md_out_path    = output_dir / "paragraphs.md"
        
        image_dir.mkdir(parents=True, exist_ok=True)
        for illustration in illustrations:
            illustration.export(image_dir) 
                    
        # with open(pages_out_path, "w", encoding="utf-8") as f:
        #     json.dump(pages, f, ensure_ascii=False, indent=2)
        
        for illustraion in illustrations:
            illustraion.export(image_dir)
        
        with open(md_out_path, "w", encoding="utf-8") as f:
            for i, para in enumerate(paragraphs):
                # Markdown output
                f.write(f"*[Paragraph: {i + 1}; Count: {para.count}]*\n\n")
                f.write(para.content)
                f.write("\n\n")
        
        print(f"📂  Result cached → {output_dir}")
    
    return Body(
        paragraphs=paragraphs,
        figures=[],
        total_words=total_count
    )


def dump_docx_body_to_md(
    body        : Body,
    md_out_path : PathLike[str],
    image_dir   : PathLike[str]
):
            
    image_dir.mkdir(parents=True, exist_ok=True)
    
    for figures in body.figures:
        figures.export(image_dir)
    
    with open(md_out_path, "w", encoding="utf-8") as f:
        for i, para in enumerate(body.paragraphs):    
            if len(para.page_ids) == 1:
                page_str = str(para.page_ids[0])
            elif len(para.page_ids) > 1:
                page_str = f'{para.page_ids[0]}–{para.page_ids[-1]}'
            else:
                page_str = ''
            
            # Markdown output
            f.write(f"*[Paragraph: {i + 1}; Count: {para.count}; Pages: {page_str}]*\n\n")
            f.write(para.content)
            f.write("\n\n")
    
    print(f"📂  Result cached → {md_out_path}")
